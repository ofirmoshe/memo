import os
import mimetypes
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
import base64
from io import BytesIO

# Document processing
try:
    import PyPDF2
    import docx
    from PIL import Image
    HAS_DOC_PROCESSING = True
except ImportError:
    HAS_DOC_PROCESSING = False
    logging.warning("Document processing libraries not installed. Install with: pip install PyPDF2 python-docx pillow")

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles processing of different file types for content extraction."""
    
    SUPPORTED_IMAGE_TYPES = {
        'image/jpeg', 'image/jpg', 'image/png', 'image/gif', 'image/bmp', 'image/webp'
    }
    
    SUPPORTED_DOCUMENT_TYPES = {
        'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword', 'text/plain', 'text/csv'
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    UPLOAD_DIR = "uploads"
    
    def __init__(self):
        # Create upload directory if it doesn't exist
        os.makedirs(self.UPLOAD_DIR, exist_ok=True)
    
    def is_supported_file_type(self, mime_type: str) -> bool:
        """Check if the file type is supported."""
        return mime_type in self.SUPPORTED_IMAGE_TYPES or mime_type in self.SUPPORTED_DOCUMENT_TYPES
    
    def get_file_category(self, mime_type: str) -> str:
        """Get the category of a file based on its MIME type."""
        if mime_type in self.SUPPORTED_IMAGE_TYPES:
            return "image"
        elif mime_type in self.SUPPORTED_DOCUMENT_TYPES:
            return "document"
        else:
            return "unknown"
    
    def save_file(self, file_data: bytes, filename: str, user_id: str) -> Tuple[str, str, int]:
        """
        Save uploaded file data to disk.
        
        Args:
            file_data: Raw file data
            filename: Original filename
            user_id: User ID for organizing files
            
        Returns:
            Tuple of (file_path, mime_type, file_size)
        """
        try:
            # Create user-specific directory
            user_dir = os.path.join(self.UPLOAD_DIR, user_id)
            os.makedirs(user_dir, exist_ok=True)
            
            # Generate safe filename
            safe_filename = self._generate_safe_filename(filename)
            file_path = os.path.join(user_dir, safe_filename)
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                # Fallback MIME type detection based on extension
                ext = Path(filename).suffix.lower()
                mime_type = self._get_mime_type_from_extension(ext)
            
            file_size = len(file_data)
            
            logger.info(f"Saved file: {file_path} ({mime_type}, {file_size} bytes)")
            return file_path, mime_type, file_size
            
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            raise
    
    def extract_content_from_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extract content from a file.
        
        Args:
            file_path: Path to the file
            mime_type: MIME type of the file
            
        Returns:
            Dict containing extracted content and metadata
        """
        try:
            if mime_type in self.SUPPORTED_IMAGE_TYPES:
                return self._extract_from_image_multimodal(file_path, mime_type)
            elif mime_type in self.SUPPORTED_DOCUMENT_TYPES:
                return self._extract_from_document(file_path, mime_type)
            else:
                return {
                    'error': f'Unsupported file type: {mime_type}',
                    'text_content': '',
                    'metadata': {}
                }
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {str(e)}")
            return {
                'error': str(e),
                'text_content': '',
                'metadata': {}
            }
    
    def _extract_from_image_multimodal(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extract content from images using multimodal LLM (no OCR needed).
        
        Note: The actual multimodal analysis will be done in the main application
        using the analyze_image_with_llm function. This method just validates
        the image and prepares metadata.
        
        Args:
            file_path: Path to the image file
            mime_type: MIME type of the image
            
        Returns:
            Dict with image metadata and placeholder for LLM analysis
        """
        try:
            # Validate image file
            with Image.open(file_path) as img:
                width, height = img.size
                format_name = img.format
                mode = img.mode
            
            metadata = {
                'width': width,
                'height': height,
                'format': format_name,
                'mode': mode,
                'mime_type': mime_type,
                'processing_method': 'multimodal_llm'
            }
            
            return {
                'text_content': '',  # Will be filled by multimodal LLM
                'metadata': metadata,
                'requires_llm_analysis': True,  # Flag to indicate LLM analysis needed
                'image_path': file_path
            }
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            return {
                'error': f'Error processing image: {str(e)}',
                'text_content': '',
                'metadata': {'mime_type': mime_type}
            }
    
    def _extract_from_document(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extract text content from document files.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            
        Returns:
            Dict containing extracted text and metadata
        """
        if not HAS_DOC_PROCESSING:
            return {
                'error': 'Document processing libraries not available',
                'text_content': '',
                'metadata': {}
            }
        
        try:
            text_content = ""
            metadata = {'mime_type': mime_type}
            
            if mime_type == 'application/pdf':
                text_content, pdf_metadata = self._extract_from_pdf(file_path)
                metadata.update(pdf_metadata)
                
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                text_content, doc_metadata = self._extract_from_docx(file_path)
                metadata.update(doc_metadata)
                
            elif mime_type in ['text/plain', 'text/csv']:
                text_content, txt_metadata = self._extract_from_text(file_path)
                metadata.update(txt_metadata)
            
            return {
                'text_content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error extracting from document {file_path}: {str(e)}")
            return {
                'error': str(e),
                'text_content': '',
                'metadata': {'mime_type': mime_type}
            }
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF file."""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
            
            metadata = {
                'num_pages': num_pages,
                'processing_method': 'PyPDF2'
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            metadata = {
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'processing_method': 'python-docx'
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            raise
    
    def _extract_from_text(self, file_path: str) -> Tuple[str, Dict]:
        """Extract content from text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            metadata = {
                'encoding': 'utf-8',
                'processing_method': 'direct_read'
            }
            
            return text_content, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                metadata = {
                    'encoding': 'latin-1',
                    'processing_method': 'direct_read'
                }
                return text_content, metadata
            except Exception as e:
                logger.error(f"Error reading text file with fallback encoding: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error extracting from text file: {str(e)}")
            raise
    
    def _generate_safe_filename(self, filename: str) -> str:
        """Generate a safe filename for storage."""
        # Remove path components and sanitize
        safe_name = os.path.basename(filename)
        # Replace potentially problematic characters
        safe_name = "".join(c for c in safe_name if c.isalnum() or c in '._-')
        
        # Add timestamp to avoid conflicts
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name, ext = os.path.splitext(safe_name)
        
        return f"{name}_{timestamp}{ext}"
    
    def _get_mime_type_from_extension(self, extension: str) -> str:
        """Get MIME type from file extension."""
        ext_to_mime = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp'
        }
        
        return ext_to_mime.get(extension.lower(), 'application/octet-stream') 